import os
import io
from pathlib import Path
import logging
import re
from collections import Counter
from typing import Optional, Union, List, Dict

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# PDF processing
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document processing
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TextExtractor:
    """Extract and clean text from various file formats"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']

    def extract_text(self, file_path: Union[str, Path, io.BytesIO],
                     file_type: Optional[str] = None) -> str:
        """
        Extract text from uploaded file

        Args:
            file_path: Path to file or BytesIO object
            file_type: File extension if using BytesIO

        Returns:
            Cleaned text content
        """
        try:
            if isinstance(file_path, io.BytesIO):
                if not file_type:
                    raise ValueError("file_type required for BytesIO objects")
                return self._extract_from_bytes(file_path, file_type)

            file_path = Path(file_path)
            file_type = file_path.suffix.lower()

            if file_type == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_type in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_type == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_type}")

        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return f"Error extracting text: {str(e)}"

    def _extract_from_bytes(self, file_bytes: io.BytesIO, file_type: str) -> str:
        """Extract text from BytesIO object"""
        file_type = file_type.lower()

        if file_type == '.pdf':
            return self._extract_pdf_from_bytes(file_bytes)
        elif file_type in ['.docx', '.doc']:
            return self._extract_docx_from_bytes(file_bytes)
        elif file_type == '.txt':
            file_bytes.seek(0)
            return file_bytes.read().decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {file_type}")

    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            return "PDF support not available. Please install PyPDF2: pip install PyPDF2"

        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            return f"Error reading PDF: {str(e)}"

        return self._clean_text(text)

    def _extract_pdf_from_bytes(self, file_bytes: io.BytesIO) -> str:
        """Extract text from PDF BytesIO"""
        if not PDF_AVAILABLE:
            return "PDF support not available. Please install PyPDF2: pip install PyPDF2"

        text = ""
        try:
            file_bytes.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_bytes)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF from bytes: {str(e)}")
            return f"Error reading PDF: {str(e)}"

        return self._clean_text(text)

    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            return "Word document support not available. Please install python-docx: pip install python-docx"

        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            return f"Error reading DOCX: {str(e)}"

    def _extract_docx_from_bytes(self, file_bytes: io.BytesIO) -> str:
        """Extract text from Word document BytesIO"""
        if not DOCX_AVAILABLE:
            return "Word document support not available. Please install python-docx: pip install python-docx"

        try:
            file_bytes.seek(0)
            doc = docx.Document(file_bytes)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error reading DOCX from bytes: {str(e)}")
            return f"Error reading DOCX: {str(e)}"

    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            return f"Error reading text file: {str(e)}"

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)

        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\-\.\,\(\)\/\@]', ' ', text)

        # Remove extra spaces
        text = ' '.join(text.split())

        return text.strip()

    def extract_resume_sections(self, text: str) -> Dict[str, str]:
        """
        Extract structured sections from resume text

        Args:
            text: Raw resume text

        Returns:
            Dictionary with resume sections
        """
        sections = {
            'contact': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'projects': '',
            'other': ''
        }

        # Common section headers (case insensitive)
        section_patterns = {
            'contact': r'(contact|personal|phone|email|address)',
            'summary': r'(summary|objective|profile|about)',
            'experience': r'(experience|work|employment|career|professional)',
            'education': r'(education|academic|school|university|degree)',
            'skills': r'(skills|technical|technologies|expertise|competencies)',
            'projects': r'(projects|portfolio|work samples|achievements)'
        }

        # Split text into potential sections
        lines = text.split('\n')
        current_section = 'other'

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if line is a section header
            found_section = False
            for section, pattern in section_patterns.items():
                if re.search(pattern, line, re.IGNORECASE):
                    current_section = section
                    found_section = True
                    break

            # Add line to current section if it's not a header
            if not found_section:
                sections[current_section] += line + ' '

        # Clean up sections
        for key in sections:
            sections[key] = sections[key].strip()

        return sections

    def extract_keywords(self, text: str) -> List[str]:
        """
        Extract important keywords from resume text

        Args:
            text: Resume text

        Returns:
            List of important keywords
        """
        # Common technical terms and skills
        tech_keywords = [
            'python', 'java', 'javascript', 'react', 'node', 'sql', 'html', 'css',
            'machine learning', 'data science', 'ai', 'artificial intelligence',
            'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'google cloud',
            'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn',
            'spring', 'django', 'flask', 'express', 'angular', 'vue',
            'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch',
            'git', 'github', 'gitlab', 'jenkins', 'ci/cd', 'devops',
            'agile', 'scrum', 'kanban', 'project management',
            'linux', 'windows', 'mac', 'unix', 'bash', 'powershell',
            'rest api', 'graphql', 'microservices', 'api',
            'bootstrap', 'tailwind', 'material-ui', 'ui/ux',
            'testing', 'junit', 'pytest', 'selenium', 'cypress',
            'data analysis', 'statistics', 'excel', 'tableau', 'power bi',
            'business intelligence', 'etl', 'data warehouse',
            'leadership', 'management', 'communication', 'teamwork',
            'problem solving', 'analytical', 'creative', 'innovative'
        ]

        # Soft skills and general terms
        soft_skills = [
            'leadership', 'management', 'communication', 'teamwork',
            'collaboration', 'problem solving', 'analytical', 'creative',
            'innovative', 'adaptable', 'flexible', 'organized',
            'detail-oriented', 'motivated', 'results-driven',
            'customer service', 'sales', 'marketing', 'business development',
            'strategic planning', 'project management', 'budget management',
            'training', 'mentoring', 'coaching', 'presentation',
            'negotiation', 'conflict resolution', 'time management'
        ]

        # Education and certifications
        education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree',
            'certification', 'certified', 'diploma', 'course',
            'training', 'workshop', 'seminar', 'conference',
            'aws certified', 'microsoft certified', 'google certified',
            'pmp', 'scrum master', 'product owner', 'safe',
            'comptia', 'cisco', 'vmware', 'oracle', 'salesforce'
        ]

        # Combine all keywords
        all_keywords = tech_keywords + soft_skills + education_keywords

        # Convert text to lowercase for matching
        text_lower = text.lower()

        # Find matching keywords
        found_keywords = []
        for keyword in all_keywords:
            if keyword.lower() in text_lower:
                found_keywords.append(keyword)

        # Extract additional keywords using regex patterns
        # Programming languages
        prog_langs = re.findall(r'\b(python|java|javascript|c\+\+|c#|php|ruby|go|rust|swift|kotlin|scala|r)\b',
                                text_lower)

        # Years of experience
        experience_years = re.findall(r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)', text_lower)

        # Degrees
        degrees = re.findall(r'\b(b\.?s\.?|m\.?s\.?|ph\.?d\.?|bachelor|master|doctorate)\b', text_lower)

        # Add found patterns to keywords
        found_keywords.extend(prog_langs)
        found_keywords.extend([f"{year} years experience" for year in experience_years])
        found_keywords.extend(degrees)

        # Remove duplicates and return
        return list(set(found_keywords))

    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """
        Extract contact information from resume text

        Args:
            text: Resume text

        Returns:
            Dictionary with contact information
        """
        contact_info = {
            'email': '',
            'phone': '',
            'linkedin': '',
            'github': '',
            'location': ''
        }

        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]

        # Phone pattern
        phone_pattern = r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = ''.join(phones[0])

        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/([a-zA-Z0-9\-]+)'
        linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            contact_info['linkedin'] = f"linkedin.com/in/{linkedin_matches[0]}"

        # GitHub pattern
        github_pattern = r'github\.com/([a-zA-Z0-9\-]+)'
        github_matches = re.findall(github_pattern, text, re.IGNORECASE)
        if github_matches:
            contact_info['github'] = f"github.com/{github_matches[0]}"

        # Location pattern (basic city, state detection)
        location_pattern = r'\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)*),\s*([A-Z]{2})\b'
        locations = re.findall(location_pattern, text)
        if locations:
            contact_info['location'] = f"{locations[0][0]}, {locations[0][1]}"

        return contact_info

    def extract_years_of_experience(self, text: str) -> int:
        """
        Extract total years of experience from resume text

        Args:
            text: Resume text

        Returns:
            Estimated years of experience
        """
        # Look for explicit mentions of years of experience
        experience_patterns = [
            r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+)\+\s*(?:years?|yrs?)',
            r'over\s*(\d+)\s*(?:years?|yrs?)',
            r'more\s*than\s*(\d+)\s*(?:years?|yrs?)'
        ]

        max_years = 0
        text_lower = text.lower()

        for pattern in experience_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                years = int(match)
                max_years = max(max_years, years)

        # If no explicit mention, try to estimate from work history
        if max_years == 0:
            # Look for date ranges in work experience
            date_patterns = [
                r'(\d{4})\s*[-–]\s*(\d{4})',
                r'(\d{4})\s*[-–]\s*present',
                r'(\d{4})\s*[-–]\s*current'
            ]

            current_year = 2024
            total_experience = 0

            for pattern in date_patterns:
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    start_year = int(match[0])
                    if len(match) > 1 and match[1].isdigit():
                        end_year = int(match[1])
                    else:
                        end_year = current_year

                    experience_duration = end_year - start_year
                    total_experience += max(0, experience_duration)

            max_years = total_experience

        return max_years

    def get_file_info(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """
        Get basic file information

        Args:
            file_path: Path to file

        Returns:
            Dictionary with file information
        """
        if isinstance(file_path, str):
            file_path = Path(file_path)

        return {
            'name': file_path.name,
            'extension': file_path.suffix,
            'size': str(file_path.stat().st_size) if file_path.exists() else 'Unknown',
            'supported': file_path.suffix.lower() in self.supported_formats
        }